import hashlib
import os
import re
import shutil
import subprocess
import tempfile
from functools import cached_property

from docx import Document
from pydub import AudioSegment
from tqdm import tqdm


def _to_kebab(text: str) -> str:
    text = re.sub(r"[^\w\s]", "", text.lower())
    return re.sub(r"[\s_]+", "-", text).strip("-")


class TextToSpeechCache:

    @staticmethod
    def get_temp_mp3_path(text: str, voice="Alex") -> str:
        h = hashlib.sha256(f"{voice}:{text}".encode()).hexdigest()[:16]
        DIR_TTS_CACHE = os.path.join(
            tempfile.gettempdir(), "docx_to_mp3", "tts_cache"
        )
        temp_mp3_path = os.path.join(DIR_TTS_CACHE, f"{h}.mp3")
        return temp_mp3_path

    @staticmethod
    def get_mp3(text: str, voice="Alex") -> str:
        text = text.strip()
        if len(text) == 0:
            raise ValueError("Text cannot be empty")

        temp_mp3_path = TextToSpeechCache.get_temp_mp3_path(text, voice=voice)

        if os.path.exists(temp_mp3_path):
            return temp_mp3_path

        temp_aiff_path = temp_mp3_path.replace(".mp3", ".aiff")

        subprocess.run(
            ["say", "-v", voice, "-o", temp_aiff_path, "--", text],
            check=True,
            capture_output=True,
        )

        AudioSegment.from_file(temp_aiff_path, format="aiff").export(
            temp_mp3_path, format="mp3"
        )
        os.unlink(temp_aiff_path)
        # print(f"{temp_mp3_path} <- {voice}: {text}")
        return temp_mp3_path


class Heading:
    def __init__(self, text: str):
        self.text = text

    def get_audio(self) -> AudioSegment:
        temp_file_path = TextToSpeechCache.get_mp3(self.text, voice="Jamie")
        return AudioSegment.from_mp3(temp_file_path)

    def n_words(self) -> int:
        return len(self.text.split())


class Paragraph:
    def __init__(self, text: str):
        self.text = text

    def get_audio(self) -> AudioSegment:
        temp_file_path = TextToSpeechCache.get_mp3(self.text, voice="Serena")
        return AudioSegment.from_mp3(temp_file_path)

    def n_words(self) -> int:
        return len(self.text.split())


class Chapter:
    T_SILENCE_BEFORE_HEADING_MS = 250
    T_SILENCE_AFTER_HEADING_MS = 500
    T_SILENCE_AFTER_PARAGRAPH_MS = 250

    def __init__(
        self, i_chapter: int, heading: Heading, paragraphs: list[Paragraph]
    ):
        self.i_chapter = i_chapter
        self.heading = heading
        self.paragraphs = paragraphs

    def get_audio(self) -> AudioSegment:
        with tqdm(
            total=self.n_words(),
            desc=f"  Chapter {self.i_chapter:02d}",
            unit="word",
            position=1,
            leave=False,
        ) as pbar:
            audio = (
                AudioSegment.silent(duration=self.T_SILENCE_BEFORE_HEADING_MS)
                + self.heading.get_audio()
                + AudioSegment.silent(duration=self.T_SILENCE_AFTER_HEADING_MS)
            )
            pbar.update(self.heading.n_words())
            for para in self.paragraphs:
                audio += para.get_audio()
                audio += AudioSegment.silent(
                    duration=self.T_SILENCE_AFTER_PARAGRAPH_MS
                )
                pbar.update(para.n_words())
        return audio

    def get_text(self) -> str:
        text = f"{self.heading.text}\n"
        for para in self.paragraphs:
            text += f"{para.text}\n"
        return text.strip()

    def n_words(self) -> int:
        return self.heading.n_words() + sum(
            para.n_words() for para in self.paragraphs
        )

    @cached_property
    def file_name(self) -> str:
        return f"{self.i_chapter:02d}-{_to_kebab(self.heading.text)}.mp3"

    def build_audio(self, output_chapters_dir) -> str:
        temp_file_path = TextToSpeechCache.get_temp_mp3_path(self.get_text())
        if not os.path.exists(temp_file_path):
            chapter_audio = self.get_audio()
            chapter_audio.export(temp_file_path, format="mp3")

        chapter_file_path = os.path.join(
            output_chapters_dir,
            self.file_name,
        )
        if not os.path.exists(chapter_file_path):
            shutil.copy(temp_file_path, chapter_file_path)
        return chapter_file_path


class DocxFile:
    def __init__(self, file_path: str):
        self.file_path = file_path

    @cached_property
    def chapters(self):
        doc = Document(self.file_path)
        chapters = []
        current_heading = None
        current_paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip() in ["", "---", "***"]:
                continue
            if para.style.name.startswith("Heading"):
                if current_heading is not None:
                    chapters.append(
                        Chapter(
                            len(chapters) + 1,
                            current_heading,
                            current_paragraphs,
                        )
                    )
                current_heading = Heading(para.text)
                current_paragraphs = []
            else:
                current_paragraphs.append(Paragraph(para.text))

        if current_heading is not None:
            chapters.append(
                Chapter(
                    len(chapters) + 1,
                    current_heading,
                    current_paragraphs,
                )
            )

        return chapters

    def n_words(self) -> int:
        return sum(chapter.n_words() for chapter in self.chapters)

    def get_audio(self) -> AudioSegment:
        audio = AudioSegment.silent(duration=0)
        for chapter in self.chapters:
            audio += chapter.get_audio()
        return audio

    def build_audio(self):
        output_dir = os.path.splitext(self.file_path)[0] + ".audio"
        shutil.rmtree(output_dir, ignore_errors=True)
        os.makedirs(output_dir, exist_ok=True)

        n_words = self.n_words()
        print(f"word-count: {n_words}")

        # chapters
        output_chapters_dir = os.path.join(output_dir, "chapters")
        os.makedirs(output_chapters_dir, exist_ok=True)

        n_chapters = len(self.chapters)
        with tqdm(
            total=n_chapters, desc="Book", unit="ch", position=0
        ) as book_pbar:
            for i_chapter, chapter in enumerate(self.chapters, start=1):
                book_pbar.set_description(
                    f"Chapter {i_chapter:02d}/{n_chapters}"
                )
                chapter_file_path = chapter.build_audio(output_chapters_dir)
                book_pbar.update(1)
                size_mb = os.path.getsize(chapter_file_path) / 1e6
                tqdm.write(f"✅ {chapter.file_name} ({size_mb:.2f} MB)")
